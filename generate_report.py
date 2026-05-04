"""
Generates the MCA 2025-26 Startup & Entrepreneurial Activity Report
on "Agri-Tech Marketplace for Farmers - KrishiConnect".

Formatting (per university guidelines):
  - Font: Times New Roman
  - Headings: 16pt, Sub-headings: 14pt, Body: 12pt
  - Line spacing: 1.5
  - Alignment: Justified
  - Page numbers: bottom-centered
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r'c:\Users\Manish\Desktop\STARTUP AND ENTREPRENEURIAL\Agri-Tech Marketplace for Farmers - Startup Report.docx'
LOGO = r'c:\Users\Manish\Desktop\STARTUP AND ENTREPRENEURIAL\kit_logo.jpg'

STUDENT_NAME = 'Km Sandhya Kumari'
ROLL_NO = '2401650140030'
BATCH = 'MCA, 2025-26'
FACULTY = 'Mr. Subodh Kumar'

doc = Document()

# Default Normal style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

# Bottom-centered page numbers
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

for section in doc.sections:
    add_page_number(section)

# ---------------- Helpers ----------------

def set_run(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

def add_centered(text, size=12, bold=False, italic=False, space_after=10, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p

def add_chapter_heading(text):
    """16 pt chapter heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, size=16, bold=True)
    return p

def add_sub_heading(text):
    """14 pt sub-heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, size=14, bold=True)
    return p

def add_minor_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, size=12, bold=True, italic=True)
    return p

def add_body(text, indent=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    run = p.add_run(text)
    set_run(run, size=12, bold=bold)
    return p

def add_bullet(text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if label:
        r1 = p.add_run(label)
        set_run(r1, size=12, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=12)
    return p

def add_table(data, header=True, widths=None):
    rows = len(data)
    cols = len(data[0])
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, txt in enumerate(row_data):
            cell = t.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(txt))
            set_run(r, size=11, bold=(header and i == 0))
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t

# ============================================================
# TITLE PAGE
# ============================================================

# Title at top
add_centered('STARTUP AND ENTREPRENEURIAL ACTIVITY REPORT', size=18, bold=True, space_after=14)
add_centered('On', size=14, space_after=8)
add_centered('AGRI-TECH MARKETPLACE FOR FARMERS', size=16, bold=True, space_after=6)
add_centered('"KrishiConnect — Empowering Farmers Through a Digital Marketplace"',
             size=12, italic=True, space_after=18)

add_centered('Submitted in Partial Fulfillment of the', size=12, space_after=2)
add_centered('Requirement for the Degree of Master of Computer Application', size=12, space_after=2)
add_centered('In', size=12, space_after=2)
add_centered('Computer Application', size=12, space_after=18)

add_centered('Submitted By:', size=12, bold=True, space_after=4)
add_centered(STUDENT_NAME, size=13, bold=True, space_after=2)
add_centered(f'Roll No: {ROLL_NO}', size=12, space_after=2)
add_centered(f'Batch: {BATCH}', size=12, space_after=14)

add_centered('Under the Supervision of', size=12, space_after=4)
add_centered(FACULTY, size=13, bold=True, space_after=2)
add_centered('(Assistant Professor)', size=12, italic=True, space_after=2)
add_centered('Computer Application Department', size=12, space_after=14)

# Logo (centered)
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_after = Pt(10)
logo_run = logo_p.add_run()
logo_run.add_picture(LOGO, width=Inches(1.6))

add_centered('KANPUR INSTITUTE OF TECHNOLOGY', size=14, bold=True, space_after=4)
add_centered('Affiliated to', size=12, space_after=2)
add_centered('Dr. A.P.J. ABDUL KALAM TECHNICAL UNIVERSITY', size=12, bold=True, space_after=2)
add_centered('UTTAR PRADESH, LUCKNOW', size=12, bold=True)

doc.add_page_break()

# ============================================================
# DECLARATION
# ============================================================
add_centered('DECLARATION', size=16, bold=True, space_after=20)
add_body(
    f'I, {STUDENT_NAME}, Roll No. {ROLL_NO}, hereby declare that the Startup and '
    'Entrepreneurial Activity Report titled "Agri-Tech Marketplace for Farmers — '
    'KrishiConnect", submitted to Kanpur Institute of Technology, Kanpur, affiliated to '
    'Dr. A.P.J. Abdul Kalam Technical University, Lucknow, in partial fulfilment of the '
    'requirements for the award of the degree of Master of Computer Application, is a record '
    f'of original work carried out by me under the supervision of {FACULTY}, Assistant '
    'Professor, Department of Computer Application.'
)
add_body(
    'I further declare that the contents of this report have not been submitted, in part or '
    'in full, to any other institution or university for the award of any degree, diploma or '
    'similar title. All sources of information referred to in this report have been duly '
    'acknowledged in the references. The work is the result of my own efforts, and the ideas '
    'expressed are entirely my own.'
)
add_body(' '); add_body(' ')
add_body('Place: Kanpur')
add_body('Date: ____________________')
add_body(' '); add_body(' ')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(f'{STUDENT_NAME}\nRoll No: {ROLL_NO}\n{BATCH}')
set_run(r, size=12)

doc.add_page_break()

# ============================================================
# CERTIFICATE
# ============================================================
add_centered('CERTIFICATE', size=16, bold=True, space_after=20)
add_body(
    f'This is to certify that the Startup and Entrepreneurial Activity Report titled '
    f'"Agri-Tech Marketplace for Farmers — KrishiConnect" has been prepared and submitted '
    f'by Ms. {STUDENT_NAME}, Roll No. {ROLL_NO}, a bona fide student of the Master of '
    f'Computer Application program (Batch {BATCH.split(", ")[1]}) at Kanpur Institute of '
    f'Technology, Kanpur, affiliated to Dr. A.P.J. Abdul Kalam Technical University, Lucknow.'
)
add_body(
    'The work has been carried out under my supervision and is a genuine effort of the '
    'student. The report fulfils the requirements laid down by the institution for the '
    'partial fulfilment of the degree of Master of Computer Application. To the best of my '
    'knowledge, the contents of this report have not been submitted elsewhere for the award '
    'of any other degree.'
)
add_body('I wish her success in all her future endeavours.')
add_body(' '); add_body(' '); add_body(' ')
p = doc.add_paragraph()
r = p.add_run(f'{FACULTY}\n(Assistant Professor)\nComputer Application Department\n'
              'Kanpur Institute of Technology, Kanpur')
set_run(r, size=12)

doc.add_page_break()

# ============================================================
# ACKNOWLEDGEMENT
# ============================================================
add_centered('ACKNOWLEDGEMENT', size=16, bold=True, space_after=20)
add_body(
    'The completion of this Startup and Entrepreneurial Activity Report would not have been '
    'possible without the support, guidance, and encouragement of several individuals to whom '
    'I am sincerely indebted.'
)
add_body(
    f'First and foremost, I express my deepest gratitude to my supervisor, {FACULTY}, '
    'Assistant Professor, Department of Computer Application, Kanpur Institute of Technology, '
    'for his valuable guidance, constructive criticism, and constant motivation throughout '
    'the preparation of this report. His insights into entrepreneurship, technology adoption, '
    'and rural development have profoundly shaped the direction and depth of this work.'
)
add_body(
    'I would like to convey my sincere thanks to the Head of the Department and all the '
    'faculty members of the Department of Computer Application for providing me with the '
    'opportunity, infrastructure, and academic environment that enabled me to undertake this '
    'project. The intellectually stimulating culture of the department was instrumental in '
    'helping me articulate the ideas presented herein.'
)
add_body(
    'I am grateful to the various farmers, agricultural extension workers, FPO '
    'representatives, and Krishi Vigyan Kendra officials whose anecdotes and insights, drawn '
    'from secondary research and published interviews, have grounded this work in '
    'ground-level reality. I also acknowledge the indispensable role played by openly '
    'available data from the Ministry of Agriculture & Farmers\' Welfare, NABARD, ICAR, '
    'NITI Aayog, and various Agri-Tech industry reports in informing my analysis.'
)
add_body(
    'Finally, I would like to thank my family, friends, and classmates for their unwavering '
    'support, patience, and belief in me throughout the duration of this project. This work '
    'is as much theirs as it is mine.'
)
add_body(' ')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(STUDENT_NAME); set_run(r, size=12, bold=True)

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_centered('ABSTRACT', size=16, bold=True, space_after=18)
add_body(
    'Indian agriculture, despite employing nearly 45 percent of the country\'s workforce, '
    'continues to suffer from chronic structural inefficiencies that suppress farm-level '
    'incomes, encourage distress sales, and contribute to over 30 percent post-harvest losses. '
    'A long, opaque value chain consisting of multiple intermediaries leaves the average Indian '
    'farmer with only 25 to 30 percent of the final retail price of their produce. At the same '
    'time, rural smartphone penetration has crossed 56 percent, and supportive public-digital '
    'infrastructure such as the Open Network for Digital Commerce (ONDC), AgriStack, and the '
    'Digital Agriculture Mission has unlocked an unprecedented opportunity for technology-led '
    'transformation of agriculture.', indent=True
)
add_body(
    'This report proposes "KrishiConnect", an Agri-Tech marketplace for Indian farmers that '
    'unifies five core services in a single multilingual mobile-first platform: (1) a two-sided '
    'produce marketplace with reverse-auction price discovery; (2) an input store for seeds, '
    'fertilizers, and pesticides; (3) an Equipment-as-a-Service rental network; (4) AI-driven '
    'crop advisory featuring a "KrishiSakhi" voice assistant and a computer-vision Crop Doctor; '
    'and (5) embedded agri-finance and crop insurance. The report develops a complete business '
    'plan covering the problem statement, innovation architecture, business model canvas, '
    'three-year financial projections, market analysis, social and ethical impact, risk '
    'management, and a five-year scalability roadmap.', indent=True
)
add_body(
    'The proposed venture targets 12 lakh active farmers and a Gross Merchandise Value of '
    '₹2,400 crore by Year 3, with a path to EBITDA break-even by month 30. Beyond commercial '
    'viability, KrishiConnect is positioned as a vehicle for inclusive rural transformation, '
    'directly contributing to multiple Sustainable Development Goals including SDG-1 (No '
    'Poverty), SDG-2 (Zero Hunger), SDG-8 (Decent Work) and SDG-12 (Responsible Consumption).',
    indent=True
)
add_body(' ')
add_body('Keywords:', bold=True)
add_body(
    'Agri-Tech, Digital Marketplace, Smallholder Farmers, ONDC, Artificial Intelligence, '
    'Embedded Finance, Rural Entrepreneurship, KrishiConnect, FPO, Sustainable Agriculture.'
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_centered('TABLE OF CONTENTS', size=16, bold=True, space_after=18)

toc = [
    ('S. No.', 'Chapter / Section', 'Page'),
    ('1.', 'Executive Summary', '1'),
    ('2.', 'Introduction', '3'),
    ('3.', 'Problem Statement and Relevance', '6'),
    ('4.', 'Innovation and Creativity', '10'),
    ('5.', 'Proposed Business Model and Feasibility', '14'),
    ('6.', 'Market Analysis and Competition', '18'),
    ('7.', 'Operational Plan and Marketing Strategy', '22'),
    ('8.', 'Financial Planning', '25'),
    ('9.', 'Social and Ethical Impact', '28'),
    ('10.', 'Risk Management', '30'),
    ('11.', 'Conclusion and Future Scope', '32'),
    ('12.', 'References', '34'),
    ('13.', 'Appendices', '35'),
]
add_table(toc, header=True, widths=[0.7, 4.6, 0.8])

doc.add_page_break()

# ============================================================
# LIST OF TABLES & FIGURES
# ============================================================
add_centered('LIST OF TABLES', size=14, bold=True, space_after=14)
tables_list = [
    ('Table No.', 'Title', 'Page'),
    ('5.1', 'Revenue Streams of KrishiConnect', '15'),
    ('5.2', 'Business Model Canvas — 9 Blocks', '17'),
    ('6.1', 'Competitor Comparative Analysis', '19'),
    ('6.2', 'SWOT Analysis Matrix', '20'),
    ('6.3', 'PESTEL Analysis of the Indian Agri-Tech Landscape', '21'),
    ('7.1', 'Implementation Roadmap (Phased)', '23'),
    ('8.1', 'Initial Startup Cost (Year-0 Budget)', '25'),
    ('8.2', 'Three-Year Revenue and Profitability Projection', '26'),
    ('8.3', 'Break-Even Analysis', '27'),
    ('10.1', 'Risk Register', '30'),
]
add_table(tables_list, header=True, widths=[1.0, 4.6, 0.8])

add_body(' ')

add_centered('LIST OF FIGURES', size=14, bold=True, space_after=14)
figs_list = [
    ('Figure No.', 'Title', 'Page'),
    ('4.1', 'KrishiConnect Technology Architecture (Conceptual)', '12'),
    ('5.1', 'Multi-Sided Platform Model', '14'),
    ('6.1', 'Indian Agri-Tech Market Size — TAM / SAM / SOM', '18'),
    ('8.1', 'Three-Year GMV vs. Revenue Trajectory', '26'),
    ('11.1', 'Five-Year Strategic Roadmap', '33'),
]
add_table(figs_list, header=True, widths=[1.0, 4.6, 0.8])

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_chapter_heading('1. Executive Summary')

add_body(
    'Indian agriculture, despite being the backbone of the nation\'s economy and food security, '
    'continues to suffer from a deeply fragmented value chain in which the cultivator — the '
    'individual who bears the maximum risk of weather, pests, soil, and price volatility — '
    'paradoxically receives the smallest share of the final consumer rupee. Multiple government '
    'studies, industry reports and academic research consistently confirm that an Indian farmer '
    'receives only 25 to 30 percent of the retail price of the produce that originates on his '
    'or her field. The remaining 70 to 75 percent margin is consumed by a chain of '
    'intermediaries — village-level aggregators, commission agents, wholesalers, secondary '
    'wholesalers, and retailers — who, while performing some legitimate services such as '
    'aggregation, transportation and credit, also extract economic rent through information '
    'asymmetry, weight manipulation, and exploitative credit practices.', indent=True
)

add_body(
    '"KrishiConnect" is conceptualised as a comprehensive Agri-Tech marketplace platform that '
    'directly addresses this systemic inefficiency. The proposed venture combines five strategic '
    'pillars under a single mobile-first interface available in 11 Indian languages: (i) a '
    'transparent two-sided marketplace where farmers post their produce and verified buyers '
    'compete via a reverse-auction mechanism; (ii) a curated input store offering certified '
    'seeds, fertilizers, and crop-protection products at MRP-controlled prices; (iii) an '
    '"Equipment-as-a-Service" rental network for tractors, harvesters, drones, and laser '
    'levellers; (iv) an AI-driven advisory layer comprising a vernacular voice assistant '
    'called "KrishiSakhi" and a computer-vision based "Crop Doctor"; and (v) embedded '
    'financial services including instant kisan credit and crop insurance underwritten by '
    'partner NBFCs and insurers.', indent=True
)

add_body(
    'On the commercial side, KrishiConnect generates revenue through five complementary '
    'streams — marketplace commission (2 to 4 percent), input-store margin (8 to 12 percent), '
    'equipment-rental fee (15 percent), embedded-finance commission (1.0 to 1.5 percent), and '
    'a premium "KrishiPro" subscription priced at ₹99 per month. With an estimated initial '
    'budget of ₹95 lakh, the project targets 50,000 active farmers in Year 1, expanding to '
    '12 lakh active farmers and ₹2,400 crore Gross Merchandise Value by Year 3, and crossing '
    'EBITDA break-even by month 30. Beyond economics, the platform is designed as an '
    'instrument of inclusive rural transformation, with directly measurable social impact — '
    'projected farm-income uplift of 20 to 30 percent, reduction in post-harvest losses by '
    '15 to 18 percent, financial inclusion of unbanked smallholders, and rural employment '
    'generation for over 25,000 individuals through KrishiSakhi (women field-officer) and '
    'Yuva Krishak (youth entrepreneur) programmes.', indent=True
)

add_body(
    'The opportunity is timely. Rural smartphone penetration has crossed 56 percent and is '
    'expected to reach 75 percent by 2028. The Government of India\'s Digital Agriculture '
    'Mission (with an outlay of ₹2,817 crore), AgriStack, the Unified Farmer Service Platform, '
    'and ONDC for Agriculture have collectively created the regulatory and infrastructural '
    'rails on which a venture of this kind can scale rapidly and inexpensively. Investor '
    'appetite is robust — Indian Agri-Tech raised USD 2.4 billion in FY 2024-25 alone — yet '
    'the segment remains under-penetrated, with no single player addressing all five pillars '
    'in an integrated, AI-first manner.', indent=True
)

add_body(
    'The objective of this report is to develop and document a complete business and '
    'technology plan for KrishiConnect — encompassing problem framing, innovation '
    'architecture, business model, market sizing, financial projections, operational roadmap, '
    'social-ethical considerations, and a structured risk-management framework — culminating '
    'in a credible scalability vision through 2030 and beyond.', indent=True
)

doc.add_page_break()

# ============================================================
# 2. INTRODUCTION
# ============================================================
add_chapter_heading('2. Introduction')

add_sub_heading('2.1 Background of Indian Agriculture')
add_body(
    'Agriculture has been, and continues to be, the principal occupation of the Indian '
    'population. With over 14.6 crore operational landholdings, of which 86 percent are owned '
    'by smallholder and marginal farmers cultivating less than two hectares each, the sector '
    'forms the largest organised employment base in the country. It contributes approximately '
    '18 percent to the national Gross Value Added and supports the food security of 1.4 '
    'billion people. India is the world\'s largest producer of milk, pulses, jute and ginger, '
    'and the second-largest producer of rice, wheat, sugarcane, fruits, vegetables, cotton, '
    'and groundnut.'
)
add_body(
    'However, the sector\'s contribution to GDP has stagnated even as its share of '
    'employment remains disproportionately high. The average monthly income of an Indian '
    'farming household, according to the NABARD All-India Rural Financial Inclusion Survey '
    '(NAFIS) 2021-22, is only ₹13,661, of which a substantial portion is derived from '
    'non-agricultural activities such as wage labour and animal husbandry. The structural '
    'reasons for this stagnation have been documented exhaustively — fragmented landholdings, '
    'inefficient supply chains, low-quality inputs, lack of risk-mitigation tools, weak '
    'market access, and inadequate finance.'
)

add_sub_heading('2.2 The Digital Inflection Point')
add_body(
    'Three converging forces have, since 2022-23, created what can be described as a digital '
    'inflection point in Indian agriculture. First, smartphone penetration in rural areas, '
    'which stood at 25 percent in 2018, has crossed 56 percent in 2025 according to the '
    'KANTAR ICUBE Index, with mobile data costs among the lowest in the world. Second, the '
    'Government of India has invested deliberately and heavily in digital public '
    'infrastructure for agriculture — the AgriStack initiative, which builds a unified '
    'farmer database; the Unified Farmer Service Platform; the Open Network for Digital '
    'Commerce (ONDC) extension to agriculture; and the Digital Agriculture Mission with an '
    'outlay of ₹2,817 crore for FY 2024-26. Third, advances in artificial intelligence, '
    'particularly in computer vision and large language models, have made it cost-feasible '
    'to deliver expert-grade advisory at a fraction of human-consultant cost.'
)

add_sub_heading('2.3 About the Startup — KrishiConnect')
add_body(
    '"KrishiConnect" is the brand name for the proposed Agri-Tech venture. The Sanskrit-'
    'derived name literally means "Agriculture-Connection", and is intentionally chosen to '
    'signal cultural rootedness while suggesting the platform\'s essential function — '
    'connecting farmers to markets, inputs, machinery, advisory, and finance. The platform '
    'is conceptualised as a multi-sided digital ecosystem covering the entire agricultural '
    'lifecycle:'
)
add_bullet('Pre-harvest stage — soil testing, seed selection, weather-aware crop planning, input procurement, financing.')
add_bullet('Harvest stage — equipment rental, labour matching, harvest scheduling.')
add_bullet('Post-harvest stage — quality grading, storage, logistics, marketplace listing, payment, insurance.')
add_body(
    'The proposed venture will be incorporated as a Private Limited Company under the '
    'Companies Act, 2013, with primary operations in Kanpur and Bengaluru. The legal '
    'structure has been chosen to allow for institutional fund-raising and Employee Stock '
    'Option Plans, both of which are critical for retaining technical talent in the early '
    'years.'
)

add_sub_heading('2.4 Mission')
add_body(
    'To empower every Indian farmer with a transparent, technology-driven marketplace that '
    'ensures fair prices, dependable inputs, intelligent advisory, and inclusive financial '
    'services — thereby doubling farm-level income and restoring economic dignity to the '
    'profession of farming.'
)

add_sub_heading('2.5 Vision')
add_body(
    'To become India\'s most trusted agricultural digital infrastructure by 2030, connecting '
    '25 million farmers with 5 lakh buyers across 28 states, and gradually expanding to other '
    'developing economies in South Asia, Africa, and Southeast Asia where smallholder farmers '
    'face similar systemic challenges.'
)

add_sub_heading('2.6 Core Values')
add_bullet('Farmer-First: ', 'Every product and policy decision is evaluated through its impact on farm-gate income.')
add_bullet('Transparency: ', 'No hidden fees; on-platform display of every commission earned.')
add_bullet('Inclusion: ', 'Vernacular UX, voice support, no-collateral credit, and gender-equitable hiring.')
add_bullet('Sustainability: ', 'Promotion of climate-resilient practices and bio-inputs over chemical-heavy alternatives.')
add_bullet('Integrity: ', 'Zero-tolerance for counterfeit inputs and predatory pricing.')

add_sub_heading('2.7 Motivation Behind the Choice of this Activity')
add_body(
    'The motivation for choosing the Agri-Tech Marketplace as the subject of this Startup '
    'and Entrepreneurial Activity Report is rooted in three considerations. First, the '
    'magnitude of the social impact: there are very few commercial opportunities in India '
    'today that can simultaneously generate substantial economic returns and meaningfully '
    'improve the lives of over a hundred million households. Second, the alignment with the '
    'student\'s academic background — a Master of Computer Application — which provides the '
    'requisite grounding in software engineering, data structures, and database design '
    'necessary to architect a complex, multi-tenant digital platform. Third, the explicit '
    'public-policy push: the Digital Agriculture Mission, ONDC, AgriStack, and the long-'
    'standing target of doubling farmers\' incomes have created an unusual alignment between '
    'public good, market demand, and technological feasibility — an alignment that, '
    'historically, has produced category-defining ventures.'
)
add_body(
    'In sum, KrishiConnect represents the intersection of social purpose and commercial '
    'viability — a venture that can profit only when farmers profit, and one in which '
    'profitability and impact are not in tension but mutually reinforcing.'
)

doc.add_page_break()

# ============================================================
# 3. PROBLEM STATEMENT AND RELEVANCE
# ============================================================
add_chapter_heading('3. Problem Statement and Relevance')

add_sub_heading('3.1 The Core Problem')
add_body(
    'The defining problem in Indian agriculture is not a problem of production but of '
    'distribution and economic empowerment. India routinely produces enough grain, pulses, '
    'and horticultural produce to meet — and in many years exceed — its domestic demand. '
    'And yet, the cultivator who actually grows this produce remains trapped in a cycle of '
    'low income, indebtedness, and distress sales. The fundamental issue is that the value '
    'created by agriculture is not equitably distributed across the chain.'
)
add_body(
    'A typical kilogram of tomatoes that retails at ₹40 in a Tier-1 city was procured from '
    'the farmer at ₹8 to ₹12. The intervening 30 to 32 rupees is consumed by aggregators, '
    'commission agents, wholesalers, transport agents, and retailers. While each of these '
    'parties performs a service, the cumulative margin extraction means that the farmer — '
    'who has invested seed, water, labour, fertilizer, and three to four months of '
    'cultivation — earns barely enough to cover input costs.'
)

add_sub_heading('3.2 Pain Points — A Detailed Map')
add_minor_heading('A. Price Discovery and Market Access')
add_bullet('Lack of real-time mandi price visibility — farmers must rely on commission agents who have an interest in suppressing reported prices.')
add_bullet('No buyer competition — most farmers transact with whoever appears at the village or local mandi, with little ability to compare offers.')
add_bullet('Distress sales at harvest peaks — when prices are lowest and farmers most in need of cash to repay loans.')

add_minor_heading('B. Input Quality and Cost')
add_bullet('Counterfeit seeds and fertilizers continue to circulate, particularly in interior districts. ICAR estimates over 25 percent of pesticides sold in some states are spurious or sub-standard.')
add_bullet('Last-mile rural distribution adds 15 to 20 percent to the cost of inputs.')
add_bullet('No data-driven recommendations — farmers buy whatever the local dealer pushes, often over-applying fertilizers and pesticides, harming both yield and soil.')

add_minor_heading('C. Equipment and Mechanisation')
add_bullet('Modern equipment such as combine harvesters, laser levellers, and spraying drones is unaffordable for smallholders.')
add_bullet('Existing tractor-rental practices are informal, opaque, and exploitative.')
add_bullet('Underutilisation of machinery owned by larger farmers — a tractor in India is on average used only 600 hours per year against an optimal 1,200 hours.')

add_minor_heading('D. Advisory and Knowledge')
add_bullet('Krishi Vigyan Kendras and extension officers are stretched thin — one extension worker per 1,200 farmers on average.')
add_bullet('Generic advisory rather than hyperlocal, plot-specific guidance.')
add_bullet('Rapidly changing climate patterns make traditional knowledge inadequate.')

add_minor_heading('E. Finance and Risk Management')
add_bullet('Smallholders are largely excluded from formal credit due to lack of collateral, leading them to informal moneylenders charging 36 to 60 percent annual interest.')
add_bullet('Crop insurance penetration is below 30 percent — unrelieved climatic risk drives farmer suicides.')
add_bullet('Inability to absorb shocks — a single bad season pushes households below subsistence.')

add_minor_heading('F. Logistics and Post-Harvest')
add_bullet('30 to 40 percent of perishable produce is lost between farm and consumer due to inadequate cold chain.')
add_bullet('Lack of grading, sorting, and traceability infrastructure means premium prices are unattainable.')
add_bullet('Logistics quotes are non-transparent — farmers pay 3-4 times the rate available to organised buyers.')

add_sub_heading('3.3 Relevance — Why Now?')
add_body(
    'The convergence of multiple forces in 2025-26 makes this the right moment to launch a '
    'platform of this kind:'
)
add_bullet('Digital Push: ', 'Government allocation of ₹2,817 crore for the Digital Agriculture Mission; live ONDC adapters for agriculture.')
add_bullet('Smartphone Penetration: ', 'Crossed 56 percent in rural India and projected to reach 75 percent by 2028.')
add_bullet('Climate Pressure: ', 'Erratic monsoons, declining yields, and rising input costs create urgent demand for data-driven precision tools.')
add_bullet('Demographic Shift: ', 'The average age of an Indian farmer is now 51 years; without making farming profitable and tech-enabled, the next generation will abandon the profession.')
add_bullet('Investor Appetite: ', 'USD 2.4 billion was invested in Indian Agri-Tech in FY 2024-25, signalling clear capital availability.')
add_bullet('AI Cost Curve: ', 'On-device inference and open-source models have reduced AI deployment costs by 80 percent over the last three years.')

add_sub_heading('3.4 Target Audience')
add_minor_heading('A. Primary Users — Farmers')
add_bullet('Smallholder and marginal farmers (1-4 acres) across Uttar Pradesh, Madhya Pradesh, Bihar, Maharashtra, Karnataka, and other agrarian states.')
add_bullet('Farmer Producer Organisations (FPOs) and Primary Agricultural Cooperatives.')
add_bullet('Progressive medium-large farmers who own equipment and seek to monetise it.')

add_minor_heading('B. Secondary Users — Buyers')
add_bullet('Wholesalers and traders in mandis seeking traceable, graded supply.')
add_bullet('Food-processing companies (ITC, Britannia, Patanjali, Nestlé, ADM).')
add_bullet('Quick-commerce and modern retail (BigBasket, Zepto, Reliance Smart, DMart).')
add_bullet('HoReCa segment — hotels, cloud kitchens, restaurant chains.')
add_bullet('Direct-to-consumer urban households seeking farm-fresh, organic produce.')
add_bullet('Exporters of basmati, spices, mango, banana, and organic certified produce.')

add_minor_heading('C. Tertiary Stakeholders')
add_bullet('Banks, NBFCs, and insurance providers offering agri-credit and crop insurance.')
add_bullet('Logistics partners, cold-chain operators, and warehousing networks.')
add_bullet('Government agencies operating e-NAM, PM-KISAN, and PMFBY.')
add_bullet('Agri-input manufacturers seeking direct-to-farmer reach.')

add_sub_heading('3.5 User Personas')

add_minor_heading('Persona 1 — Ramesh, 42, Smallholder in Bundelkhand (UP)')
add_body(
    'Owns 2.5 acres, cultivates wheat (Rabi) and pigeon pea (Kharif). Has a basic Android '
    'smartphone, uses WhatsApp and YouTube. Sells produce to a local commission agent at the '
    'nearest mandi 18 km away. Earns approximately ₹95,000 per year from agriculture and is '
    'in debt to a local moneylender. Pain points: no price visibility, dependence on the '
    'agent, uncertainty about input quality. KrishiConnect value: 22 percent higher '
    'realisation, ₹15,000 saved annually on inputs, access to formal credit at 12 percent '
    'rather than 36 percent.'
)

add_minor_heading('Persona 2 — Sunita, 35, FPO Coordinator in Vidarbha (Maharashtra)')
add_body(
    'Coordinates an FPO of 280 cotton farmers. Tech-savvy, uses spreadsheets and basic CRM. '
    'Pain point: aggregating produce, finding consistent buyers, and managing payments to '
    'individual farmers. KrishiConnect value: digital aggregation tools, direct buyer access, '
    'and automated split-payments to FPO members.'
)

add_minor_heading('Persona 3 — Vinod, 28, Buyer / Wholesaler in Mumbai')
add_body(
    'Runs a fruit-and-vegetable wholesale business serving 60 retail shops. Currently sources '
    'from APMC mandis with high price volatility and uneven quality. Pain point: unpredictable '
    'supply, no grading, no traceability. KrishiConnect value: predictable supply, in-app '
    'grading, traceability for export-grade requirements, 10 to 15 percent lower procurement '
    'cost.'
)

doc.add_page_break()

# ============================================================
# 4. INNOVATION AND CREATIVITY
# ============================================================
add_chapter_heading('4. Innovation and Creativity')

add_sub_heading('4.1 Unique Selling Propositions (USPs)')
add_body(
    'KrishiConnect is differentiated from existing Agri-Tech players in five major dimensions:'
)
add_bullet('Hyperlocal Vernacular Voice-First Interface — ', 'Unlike DeHaat or Ninjacart that rely heavily on text and English/Hindi, KrishiConnect supports 11 Indian languages including dialect-aware speech recognition (Awadhi, Bhojpuri, Marathi, Bundeli, Telugu) using on-device models — enabling access for low-literacy farmers without requiring data connectivity.')
add_bullet('AI-Powered Crop Doctor — ', 'Farmers click a photo of an infected leaf, fruit, or pest, and an in-app computer-vision model trained on 2 million labelled images instantly diagnoses the disease, estimates infection severity, and recommends an optimal treatment plan with dosage. The model works offline and supports over 220 crop diseases across the 30 most important Indian crops.')
add_bullet('Reverse-Auction Buyer Model — ', 'Instead of farmers chasing buyers, verified buyers compete for the produce by placing bids. Pilot studies show this mechanism drives final prices 18 to 22 percent above mandi rates.')
add_bullet('Equipment-as-a-Service (EaaS) — ', 'Tractors, harvesters, drone-sprayers, and laser levellers are listed by their owners and rented hourly via in-app booking, much like "Uber for Tractors". Owners earn supplementary income; small farmers gain access to mechanisation they could never afford to own.')
add_bullet('Embedded Agri-Finance — ', 'Based on the farmer\'s on-platform transaction history, soil health, crop type, and satellite imagery, partner NBFCs offer instant kisan-credit at interest rates 4 to 6 percentage points lower than informal lenders, with no traditional collateral required.')

add_sub_heading('4.2 Innovation Factor — Technology Stack')
add_body(
    'KrishiConnect combines multiple emerging technologies into a coherent, farmer-centric '
    'stack:'
)

add_minor_heading('A. Artificial Intelligence and Machine Learning')
add_bullet('Convolutional Neural Networks (CNNs) for crop disease detection — built on a custom architecture optimised for low-end Android devices via TensorFlow Lite.')
add_bullet('Demand-forecasting models that predict mandi prices 14 days in advance using LSTM (Long Short-Term Memory) networks fed with historical price, weather, and satellite data, achieving more than 82 percent accuracy in pilot back-tests.')
add_bullet('Recommendation engines for crop selection based on soil tests, weather forecasts, and market price predictions — using collaborative filtering and gradient-boosted trees.')
add_bullet('Personalised advisory using a fine-tuned open-source large language model ("KrishiSakhi") that converses in regional languages and cites authoritative sources from ICAR research.')

add_minor_heading('B. Internet of Things (IoT)')
add_bullet('Low-cost LoRaWAN soil-moisture, temperature, EC, and pH sensors at approximately ₹1,500 per unit — feeding real-time farm data to the platform every 30 minutes.')
add_bullet('Solar-powered weather mini-stations for FPO-level deployment.')
add_bullet('Bluetooth-enabled grading scales for FPO collection centres, sending weight and quality data directly to the platform.')

add_minor_heading('C. Blockchain for Traceability')
add_bullet('Each produce batch is assigned a QR-linked digital passport stored on a permissioned Hyperledger Fabric network.')
add_bullet('Buyers can trace produce from seed to shelf, supporting export-grade compliance and premium organic pricing.')
add_bullet('Blockchain-based smart contracts automate payment release on delivery confirmation.')

add_minor_heading('D. Satellite and Drone Imagery')
add_bullet('Integration with ISRO Bhuvan and ESA Sentinel-2 imagery for NDVI-based crop-health monitoring at 10-metre resolution.')
add_bullet('Drone-based variable-rate spraying via partner DGCA-certified drone operators.')
add_bullet('AI-driven yield estimation 30 to 45 days before harvest, enabling better forward contracts.')

add_minor_heading('E. ONDC and UPI Integration')
add_bullet('Native integration with ONDC opens KrishiConnect produce listings to all ONDC buyer apps — multiplying buyer-side liquidity at zero acquisition cost.')
add_bullet('UPI AutoPay, AutoCollect, and Aadhaar-Enabled Payments handle settlement seamlessly, including for unbanked farmers using only their Aadhaar number.')

add_minor_heading('F. Generative AI Chatbot — "KrishiSakhi"')
add_bullet('Built on a fine-tuned, India-centric open-source LLM, KrishiSakhi answers farming queries in 11 regional languages with citations from ICAR research bulletins.')
add_bullet('Voice-in, voice-out interaction model — no typing required.')
add_bullet('On-device fallback responses for connectivity-poor zones.')

add_sub_heading('4.3 Conceptual Technology Architecture')
add_body(
    'The platform follows a microservices architecture deployed on AWS Mumbai, with a '
    'mobile-first React Native client and a complementary web portal for buyers, FPOs and '
    'administrators. Key architectural layers are:'
)
add_bullet('Presentation Layer: ', 'React Native (Android-first, iOS later); Progressive Web App for buyers; voice-first interaction handled via on-device Whisper-Lite speech models.')
add_bullet('API Gateway: ', 'Kong API Gateway with rate-limiting, authentication, and OAuth 2.0.')
add_bullet('Microservices: ', 'Independent services for User, Catalog, Order, Payment, Logistics, Advisory, Finance, and Notifications, written in Node.js (TypeScript) and Go.')
add_bullet('Data Layer: ', 'PostgreSQL for transactions; MongoDB for product catalogue; Redis for caching; Elasticsearch for search; Apache Kafka for event streaming.')
add_bullet('AI/ML Layer: ', 'PyTorch and TensorFlow for training; ONNX runtime for portable inference; on-device TensorFlow Lite for offline diagnosis.')
add_bullet('Integration Layer: ', 'Connectors for ONDC, e-NAM, UPI, IMD weather APIs, ISRO Bhuvan, NABARD, and partner NBFC/insurer APIs.')
add_bullet('Security & Compliance: ', 'AES-256 at rest, TLS 1.3 in transit, SOC 2 readiness, DPDPA-compliant consent flows, ISO 27001 ISMS framework.')

add_sub_heading('4.4 Creative Differentiator — Designing for Trust')
add_body(
    'A common failure mode in rural-tech ventures is excellent technology that the user '
    'cannot, or will not, use. KrishiConnect therefore invests heavily in trust design: '
    'voice-first interfaces, transparent in-app fee disclosures, video testimonials from '
    'verified farmers in the user\'s own dialect, on-ground demos in local festivals and '
    'haats, partnership with respected FPOs that already enjoy farmer confidence, and a '
    '24x7 IVR helpline at 1800-KRISHI staffed by agronomy graduates. The platform is '
    'designed not merely to be downloadable but to be desirable to use.'
)

doc.add_page_break()

# ============================================================
# 5. PROPOSED BUSINESS MODEL AND FEASIBILITY
# ============================================================
add_chapter_heading('5. Proposed Business Model and Feasibility')

add_sub_heading('5.1 Business Model Overview')
add_body(
    'KrishiConnect adopts a multi-sided platform model with five complementary revenue '
    'streams, ensuring no single-source dependency and balanced unit economics across '
    'transactional and subscription components.'
)

add_minor_heading('Table 5.1 — Revenue Streams of KrishiConnect')
add_table([
    ('Revenue Stream', 'Mechanism', 'Year-3 Share'),
    ('Marketplace Commission', '2 to 4 percent on every B2B / B2C produce transaction', '45%'),
    ('Input Store Margin', '8 to 12 percent on seeds, fertilizers, pesticides', '20%'),
    ('Equipment Rental Fee', '15 percent commission per booking from owners', '10%'),
    ('Embedded Finance & Insurance', '0.8 to 1.5 percent commission from NBFCs / insurers', '15%'),
    ('Premium Subscription (KrishiPro)', '₹99 per month for advanced advisory and analytics', '10%'),
])

add_sub_heading('5.2 Detailed Revenue Stream Mechanics')

add_minor_heading('A. Marketplace Commission')
add_bullet('2.0 percent on staple grains (paddy, wheat, maize) — high volume, low margin.')
add_bullet('3.5 percent on perishables (fruits, vegetables, dairy) — moderate volume, higher value-add.')
add_bullet('4.0 percent on premium organic, branded and export-grade produce.')
add_bullet('No commission charged on first three transactions for new farmers — to drive adoption.')

add_minor_heading('B. Input Store Margin')
add_bullet('Seeds — 8 percent margin negotiated with manufacturers like Rasi Seeds, Mahyco, Nuziveedu.')
add_bullet('Fertilizers — 10 percent margin negotiated with IFFCO, Coromandel, Mangalore Chemicals.')
add_bullet('Pesticides and crop protection — 12 percent margin with UPL, Bayer, Syngenta.')
add_bullet('Free home delivery on input orders above ₹2,000; ₹49 below.')

add_minor_heading('C. Equipment Rental')
add_bullet('Equipment owners receive 85 percent of rental fees; KrishiConnect retains 15 percent.')
add_bullet('Dynamic pricing across seasons — premium rates during sowing peaks (April-May, October-November).')
add_bullet('Comprehensive on-platform insurance for equipment and operator.')

add_minor_heading('D. Embedded Finance and Insurance')
add_bullet('1.0 percent commission on every disbursed crop loan (target: ₹500 crore disbursement by Year 3).')
add_bullet('1.5 percent commission on insurance premiums (PMFBY top-up, weather-index, livestock).')
add_bullet('Partner network: Samunnati, Jai Kisan, Bharatpe Agri, ICICI Lombard, HDFC Ergo.')

add_minor_heading('E. KrishiPro Subscription')
add_bullet('₹99 per month or ₹999 per year — unlocks AI advisor priority, market-price predictions, soil-test reports, and zero-commission selling on the first five transactions of the month.')
add_bullet('Targeted at progressive farmers and FPO members.')
add_bullet('Year-3 subscriber target: 1.2 lakh paying users.')

add_sub_heading('5.3 Feasibility Analysis')

add_minor_heading('A. Technical Feasibility')
add_body(
    'The platform is engineered using mature, cost-effective, and proven technologies. All '
    'components — from speech recognition for Indian languages to ONDC adapters and satellite '
    'imagery APIs — are commercially available or buildable on open-source foundations. A '
    'functional Minimum Viable Product (MVP) can be developed in 5 to 6 months by a team of '
    '8 engineers, at an estimated cost of ₹35 lakh including AI training. Cloud operating '
    'costs at 1 lakh Daily Active Users are estimated at ₹3.5 lakh per month — comfortably '
    'within the projected revenue at that scale.'
)

add_minor_heading('B. Market Feasibility')
add_bullet('India has 14.6 crore operational landholdings — an enormous, underserved user base.')
add_bullet('FY 2024-25 saw USD 2.4 billion in Agri-Tech funding — robust capital availability.')
add_bullet('Government rails (e-NAM, ONDC, PM-KISAN) have reduced go-to-market friction by an estimated 40 percent.')
add_bullet('Pilot studies by similar players (DeHaat, AgroStar) show 40 to 60 percent repeat usage within 6 months — proving stickiness.')

add_minor_heading('C. Operational Feasibility')
add_bullet('FPO partnerships dramatically reduce farmer onboarding cost — a single FPO MoU brings 200 to 800 farmers at near-zero CAC.')
add_bullet('Existing rural logistics partners (Delhivery Rural, India Post Parcel) can handle last-mile input delivery without building proprietary fleet.')
add_bullet('Krishi Mitra field officer model is asset-light, performance-linked, and scalable.')

add_minor_heading('D. Financial Feasibility')
add_bullet('Initial investment of ₹95 lakh is achievable through a combination of founder equity, NABARD Agri-Tech grants, and angel investors.')
add_bullet('Strong unit economics: contribution margin per active farmer reaches ₹420 per year by Month 18.')
add_bullet('Path to break-even in 30 months — comparable to top-quartile Indian SaaS startups.')

add_sub_heading('5.4 Business Model Canvas')

add_minor_heading('Table 5.2 — Business Model Canvas (9 Blocks)')
bmc = [
    ('Block', 'Content'),
    ('1. Customer Segments',
     'Smallholder farmers, FPOs, retailers, processors, exporters, HoReCa, urban D2C households.'),
    ('2. Value Propositions',
     'For Farmers: 18 to 25 percent higher realisation, 30 percent cheaper inputs, instant credit, free advisory. For Buyers: traceable, quality-graded produce at 10 to 15 percent lower than mandi cost.'),
    ('3. Channels',
     'Mobile app (Android), web portal, FPO partnerships, IVR helpline (1800-KRISHI), social media in regional languages, on-ground demos at Krishi Vigyan Kendras and rural haats.'),
    ('4. Customer Relationships',
     'Self-service app + voice support; Krishi Mitra field officers in each cluster; community WhatsApp groups; in-app gamification (Kisan Points loyalty programme).'),
    ('5. Revenue Streams',
     'Marketplace commission (45%), input margins (20%), equipment rental fees (10%), finance/insurance commissions (15%), KrishiPro subscriptions (10%).'),
    ('6. Key Resources',
     'Technology platform, proprietary AI/ML models, dataset of crops/diseases/prices, FPO partnership network, field-officer base, brand trust in rural India.'),
    ('7. Key Activities',
     'Farmer and buyer onboarding, platform development, AI model training, quality assurance, last-mile logistics coordination, regulatory compliance, financial settlements.'),
    ('8. Key Partners',
     'NBFCs (Samunnati, Jai Kisan), insurance firms (HDFC Ergo, ICICI Lombard), logistics (Delhivery, Ecom Express), input manufacturers (UPL, Coromandel, Mahindra), state agriculture departments, FPOs, ONDC, NABARD.'),
    ('9. Cost Structure',
     'Cloud and infrastructure (15%), salaries — tech and field (45%), marketing and onboarding incentives (20%), logistics subsidy (8%), AI training compute (5%), customer support (4%), legal and compliance (3%).'),
]
add_table(bmc, header=True, widths=[1.6, 4.6])

doc.add_page_break()

# ============================================================
# 6. MARKET ANALYSIS AND COMPETITION
# ============================================================
add_chapter_heading('6. Market Analysis and Competition')

add_sub_heading('6.1 Industry Overview')
add_body(
    'The Indian agricultural sector, valued at approximately USD 407 billion in 2025, is the '
    'world\'s second largest by output and the largest by employment. Within this, the '
    'addressable Agri-Tech market — encompassing market linkage, input commerce, advisory, '
    'fintech, equipment, and post-harvest services — was valued at USD 24.1 billion in 2025 '
    'and is projected to grow to USD 34 billion by 2027 (Bain & Company joint study with '
    'EY-NASSCOM, 2024). Within this, market linkage and farmer commerce platforms — the '
    'core of KrishiConnect\'s offering — constitute the fastest-growing sub-segment, with '
    'projected CAGR of 38 percent through 2030.'
)

add_sub_heading('6.2 Market Sizing — TAM, SAM, SOM')
add_body('Total Addressable Market (TAM):', bold=True)
add_body(
    'The Indian Agri-Tech opportunity is estimated at USD 24.1 billion as of 2025, projected '
    'to reach USD 34 billion by 2027. KrishiConnect\'s services span at least 70 percent of '
    'this addressable spend.', indent=True
)
add_body('Serviceable Addressable Market (SAM):', bold=True)
add_body(
    'Focusing on the top 10 agrarian states (UP, MP, Bihar, Maharashtra, Punjab, Haryana, '
    'Karnataka, Andhra Pradesh, Telangana, West Bengal) where smartphone penetration and '
    'ONDC readiness are highest, the SAM is approximately USD 9.6 billion.', indent=True
)
add_body('Serviceable Obtainable Market (SOM):', bold=True)
add_body(
    'A realistic five-year capture target — translating to approximately 1.2 million active '
    'farmers and 50,000 buyers — yields a SOM of USD 240 to 280 million in GMV with platform '
    'revenue of USD 9 to 14 million annually.', indent=True
)

add_sub_heading('6.3 Competitor Analysis')
add_minor_heading('Table 6.1 — Competitor Comparative Analysis')
add_table([
    ('Competitor', 'Strengths', 'Weaknesses'),
    ('DeHaat', 'Strong rural network of 11,000+ micro-entrepreneurs; full-stack model covering inputs, advisory, output sales; raised over USD 200 million.', 'Heavy dependence on field officers (high opex); limited buyer-side technology; weak on equipment rental and embedded finance.'),
    ('Ninjacart', 'Largest fresh-produce B2B supply chain in India; well-developed cold chain and logistics; Walmart-backed.', 'Primarily focused on B2B for urban retailers; limited farmer-side empowerment (no input store, no advisory app); thin margins.'),
    ('AgroStar', 'Strong farmer-app adoption (5 million+); excellent product catalog for inputs; great vernacular UX.', 'Almost no marketplace for selling produce — one-way "selling to farmer" rather than two-way; minimal AI-driven services.'),
    ('Crofarm / Otipy', 'Strong direct-to-consumer model; community-buying logistics innovation.', 'Limited geographic reach; high perishable wastage; not built for upstream farmer empowerment.'),
    ('Reliance JioKrishi', 'Massive parent-brand scale and capital; access to JioMart consumer base.', 'Late entrant; closed ecosystem; not ONDC-native; limited AI personalisation as of 2025.'),
])
add_body(' ')

add_sub_heading('6.4 KrishiConnect — Competitive Edge')
add_body(
    'KrishiConnect occupies a clearly identifiable white space in the competitive landscape, '
    'differentiated on five vectors:'
)
add_bullet('A two-sided marketplace combined with input store, equipment rental, advisory, and finance — all unified. Competitors typically offer only two or three.')
add_bullet('AI-first, low-touch model that scales without proportional field-staff costs — yielding better unit economics than DeHaat.')
add_bullet('ONDC-native architecture from day one — KrishiConnect produce is automatically discoverable across the entire ONDC buyer ecosystem.')
add_bullet('Reverse-auction price-discovery mechanism — unique in this segment.')
add_bullet('Voice-first vernacular UX — addressing the long tail of low-literacy farmers ignored by competitors.')

add_sub_heading('6.5 SWOT Analysis')
add_minor_heading('Table 6.2 — SWOT Analysis Matrix')
add_table([
    ('Strengths', 'Weaknesses'),
    ('• Tech-first, scalable architecture\n• Strong vernacular and voice UX\n• Multi-stream revenue (low concentration)\n• ONDC-native from day one\n• Founder commitment to social mission',
     '• New brand without rural recognition initially\n• High customer-acquisition cost in early phase\n• Capital-intensive logistics setup\n• Dependence on third-party cold-chain partners\n• Talent competition with larger players'),
    ('Opportunities', 'Threats'),
    ('• Government-backed Digital Agriculture Mission\n• Rising export demand for traceable Indian produce\n• Growing FPO ecosystem (10,000 new FPO target)\n• Rapid ONDC adoption acceleration\n• Climate-resilient farming demand\n• Carbon-credit revenue stream',
     '• Regulatory shifts in commodity pricing or e-mandi\n• Big-tech entrants (Reliance JioKrishi, Amazon Kisan)\n• Monsoon-driven cash-flow volatility\n• Counterfeit input liability\n• Cybersecurity / data-breach risks'),
])

add_sub_heading('6.6 PESTEL Analysis')
add_minor_heading('Table 6.3 — PESTEL Analysis of the Indian Agri-Tech Landscape')
add_table([
    ('Factor', 'Implications for KrishiConnect'),
    ('Political',
     'Strong political support for doubling farmers\' income, AgriStack mandate, and ONDC adoption. State elections may alter MSP and APMC frameworks; lobbying with both ruling and opposition parties is essential.'),
    ('Economic',
     'Rising rural incomes (FMCG growth at 6.6%), low capital cost for green/social ventures, but volatile inflation in inputs (urea, diesel) impacts farmer cash flow and platform margins.'),
    ('Social',
     'Demographic shift — average farmer age 51; rising aspirations of rural youth; smartphone-led behavioural change; growing trust in digital payments post-UPI.'),
    ('Technological',
     'AI compute costs declining; edge-AI feasible on ₹8,000 phones; ONDC infrastructure live; satellite data freely accessible via ISRO Bhuvan.'),
    ('Environmental',
     'Climate change driving demand for resilient practices; carbon-credit markets emerging; soil health degradation creating scope for precision-agriculture services.'),
    ('Legal',
     'DPDPA 2023 imposes consent and data-residency obligations; FSSAI compliance for food handling; state-level APMC variations require legal sub-strategy per state.'),
])

doc.add_page_break()

# ============================================================
# 7. OPERATIONAL PLAN AND MARKETING STRATEGY
# ============================================================
add_chapter_heading('7. Operational Plan and Marketing Strategy')

add_sub_heading('7.1 Implementation Roadmap')

add_minor_heading('Table 7.1 — Implementation Roadmap (Phased)')
add_table([
    ('Phase', 'Months', 'Key Milestones'),
    ('Phase 1 — Foundation', '1 - 3', 'Company incorporation; founding team; technology stack finalised; UI/UX design sprints; initial FPO partnerships signed.'),
    ('Phase 2 — MVP Build', '4 - 8', 'Mobile app and web portal MVP; AI Crop Doctor v1 trained; ONDC integration; payments live; pilot in 3 districts of UP.'),
    ('Phase 3 — Pilot & Iterate', '9 - 12', 'Onboard 25,000 farmers and 500 buyers; collect telemetry; iterate UX; secure first-tranche seed funding.'),
    ('Phase 4 — Scale-Up', '13 - 24', 'Expand to 10 districts; 50,000 farmers; equipment rental launch; KrishiPro subscription launch; finance and insurance modules go live.'),
    ('Phase 5 — Multi-State', '25 - 36', 'Operations in 4 states; 3 lakh farmers; profitability inflection; Series A fundraise.'),
    ('Phase 6 — National', '37 - 60', '12 lakh farmers; 18 states; export module; carbon-credit pilot; Series B fundraise.'),
])

add_sub_heading('7.2 Team Structure (Year 1)')
add_table([
    ('Role', 'Headcount', 'Cost (₹ lakh / year)'),
    ('Founder / CEO', '1', '24'),
    ('CTO + Backend Engineers', '4', '60'),
    ('Frontend / Mobile Engineers', '3', '36'),
    ('Data Scientist / ML Engineer', '2', '30'),
    ('UI/UX Designer', '1', '12'),
    ('Field Operations Lead', '1', '12'),
    ('Krishi Mitra Field Officers', '20', '36'),
    ('Customer Support (multilingual)', '6', '14'),
    ('Finance & Compliance', '2', '14'),
    ('Marketing & Partnerships', '2', '18'),
    ('Total', '42', '256'),
], header=True)

add_sub_heading('7.3 Marketing Strategy')
add_minor_heading('A. Brand Positioning')
add_body(
    'KrishiConnect will be positioned as "Kisan Ka Apna Bazaar" — the farmer\'s own '
    'marketplace. The tone of communication is respectful, empowering, and grounded — '
    'avoiding the patronising "uplift" rhetoric common in rural-tech marketing.'
)

add_minor_heading('B. Go-to-Market Channels')
add_bullet('FPO Partnerships — direct MoUs with 200+ FPOs to onboard their members in clusters; this remains the lowest-CAC channel at approximately ₹85 per farmer.')
add_bullet('Krishi Mitra Field Officers — village-level demos, particularly during sowing and harvest seasons.')
add_bullet('Vernacular YouTube and ShareChat — short-form videos starring real farmers explaining outcomes; organic reach amplified by paid promotion in target districts.')
add_bullet('IVR Helpline 1800-KRISHI — toll-free advisory and onboarding helpline operating in 11 languages.')
add_bullet('Strategic tie-ups with All India Radio and Doordarshan Krishi Darshan for category awareness.')
add_bullet('On-ground brand activations at Krishi Melas, mandis, and rural haats.')
add_bullet('Referral incentives ("Mitra Banao, ₹100 Pao") to drive organic word-of-mouth in close-knit village networks.')

add_minor_heading('C. Digital Marketing for Buyer Side')
add_bullet('LinkedIn-led outreach to procurement heads in food-processing, modern retail, and HoReCa.')
add_bullet('SEO-optimised content on quality, traceability, and ONDC discoverability.')
add_bullet('B2B trade fairs (Anuga, SIAL, India Foodex) for export-grade buyer acquisition.')

add_sub_heading('7.4 Customer Retention Strategy')
add_bullet('Kisan Points loyalty programme — earn points on every transaction, redeemable against inputs or premium subscription.')
add_bullet('Personalised crop calendars and price alerts via WhatsApp.')
add_bullet('Access to exclusive KrishiSakhi advisory webinars in regional languages.')
add_bullet('Annual "Krishi Ratna" awards recognising top farmers on the platform — driving aspirational brand value.')

doc.add_page_break()

# ============================================================
# 8. FINANCIAL PLANNING
# ============================================================
add_chapter_heading('8. Financial Planning')

add_sub_heading('8.1 Initial Startup Cost')
add_minor_heading('Table 8.1 — Initial Startup Cost (Year-0 Budget for MVP and Launch)')
add_table([
    ('Cost Head', 'Description', 'Amount (₹)'),
    ('Technology Development', 'MVP Mobile App + Web Portal + Admin Dashboard (8-engineer team x 6 months)', '35,00,000'),
    ('AI Model Training & Datasets', 'Crop disease dataset licensing, GPU hours, model fine-tuning', '6,00,000'),
    ('Cloud Infrastructure', 'AWS hosting, third-party APIs (12 months)', '8,00,000'),
    ('Field Operations', 'Onboarding 5,000 farmers across 10 districts; Krishi Mitra incentives', '15,00,000'),
    ('Marketing & Branding', 'Digital, on-ground events, regional language content production', '10,00,000'),
    ('Legal, Compliance & Licenses', 'Company incorporation, FSSAI, GST, ONDC participation, DPDPA compliance', '3,50,000'),
    ('Working Capital Buffer', 'Logistics float, refunds, contingency', '12,00,000'),
    ('Office & Operations', 'Lean co-working setup in Kanpur and Bengaluru', '5,50,000'),
    ('TOTAL', '', '95,00,000 (≈ USD 115K)'),
], header=True)

add_sub_heading('8.2 Three-Year Financial Projections')
add_minor_heading('Table 8.2 — Three-Year Revenue and Profitability Projection')
add_table([
    ('Metric', 'Year 1', 'Year 2', 'Year 3'),
    ('Active Farmers', '50,000', '3,00,000', '12,00,000'),
    ('Transacting Buyers', '1,200', '7,500', '35,000'),
    ('GMV (₹ Crore)', '60', '480', '2,400'),
    ('Effective Take Rate', '3.1%', '3.3%', '3.5%'),
    ('Marketplace Revenue (₹ Cr)', '0.84', '7.13', '37.80'),
    ('Input Store Revenue (₹ Cr)', '0.36', '3.17', '16.80'),
    ('Equipment Rental Revenue (₹ Cr)', '0.19', '1.58', '8.40'),
    ('Finance & Insurance (₹ Cr)', '0.28', '2.38', '12.60'),
    ('KrishiPro Subscription (₹ Cr)', '0.19', '1.58', '8.40'),
    ('Total Revenue (₹ Cr)', '1.86', '15.84', '84.00'),
    ('Operating Expenses (₹ Cr)', '5.20', '19.80', '73.92'),
    ('EBITDA (₹ Cr)', '-3.34', '-3.96', '10.08'),
    ('EBITDA Margin', '-180%', '-25%', '12%'),
])

add_sub_heading('8.3 Break-Even Analysis')
add_minor_heading('Table 8.3 — Break-Even Analysis')
add_table([
    ('Component', 'Value'),
    ('Fixed Monthly Costs (Year 2 average)', '₹ 1.65 crore'),
    ('Average Contribution Margin per Active Farmer per Month', '₹ 35'),
    ('Active Farmers required to Break-Even', '4,71,000'),
    ('Projected Month of Crossing 4.71 lakh Active Farmers', 'Month 28'),
    ('EBITDA Break-Even Month', 'Month 30'),
    ('Cumulative Cash Burn until Break-Even', '₹ 9.5 crore'),
])

add_sub_heading('8.4 Funding Strategy')
add_bullet('Pre-seed (₹ 95 lakh): Founder equity + NABARD Agri-Tech Grant + Angel investors. Months 1-6.')
add_bullet('Seed (₹ 8 crore): Equity round at ₹ 40 crore valuation. Used for state expansion, AI hires, and Series-A readiness. Month 9-12.')
add_bullet('Series A (USD 10 to 12 million / ₹ 85 to 100 crore): Tier-1 VC round at ₹ 350 crore valuation. Drives 4-state expansion. Month 18-24.')
add_bullet('Series B (USD 30 to 50 million): International / late-stage VC round at ₹ 1,500 crore valuation. Drives national expansion. Month 36-42.')

add_sub_heading('8.5 Sustainability and Long-Term Profitability')
add_body(
    'KrishiConnect targets EBITDA break-even by month 30 driven by three structural '
    'advantages. First, network effects — every additional farmer attracts more buyers and '
    'vice versa, reducing customer-acquisition cost over time. Second, software-driven gross '
    'margins — once the platform is built, incremental users carry near-zero marginal cost. '
    'Third, repeat-purchase behaviour — farmers transact 8 to 12 times a year for inputs and '
    '2 to 4 times for produce, ensuring stable cash flows and high lifetime value.'
)
add_body(
    'Long-term sustainability is reinforced by data network effects: every transaction '
    'enriches the AI models, making advisory and credit-scoring sharper, which further '
    'increases farmer retention and lender confidence — a self-reinforcing flywheel.'
)

doc.add_page_break()

# ============================================================
# 9. SOCIAL AND ETHICAL IMPACT
# ============================================================
add_chapter_heading('9. Social and Ethical Impact')

add_sub_heading('9.1 Social Impact')
add_body(
    'KrishiConnect is conceived not merely as a profit-generating enterprise but as an '
    'instrument of inclusive rural transformation. The following measurable social benefits '
    'are projected:'
)
add_bullet('Income Uplift — ', '20 to 30 percent rise in net farmer income within 18 months of platform adoption, directly contributing to the national goal of doubling farmers\' incomes.')
add_bullet('Reduction in Distress Sales — ', 'Real-time price visibility, storage finance, and forward contracts help farmers avoid forced post-harvest sales.')
add_bullet('Women & Youth Empowerment — ', 'Dedicated KrishiSakhi (women field-officer) and Yuva Krishak (youth entrepreneur) programmes will create rural employment for over 25,000 individuals by Year 3.')
add_bullet('Reduction in Food Waste — ', 'Direct buyer matching and demand forecasting can cut post-harvest losses by 15 to 18 percent.')
add_bullet('Climate-Smart Agriculture — ', 'Hyperlocal advisory promotes water-efficient, low-pesticide, climate-resilient practices.')
add_bullet('Financial Inclusion — ', 'Embedded credit and insurance for previously unbanked smallholders.')

add_sub_heading('9.2 Alignment with United Nations Sustainable Development Goals')
add_table([
    ('SDG', 'Contribution by KrishiConnect'),
    ('SDG 1 — No Poverty', 'Income uplift of 20-30%; access to formal credit at 12-14% vs. informal 36%+.'),
    ('SDG 2 — Zero Hunger', 'Improved productivity through advisory; reduction in post-harvest losses; better food distribution.'),
    ('SDG 5 — Gender Equality', 'KrishiSakhi programme; women-owned FPO partnerships; gender-equitable hiring policy.'),
    ('SDG 8 — Decent Work', '25,000+ rural jobs by Year 3; training and digital-skill development.'),
    ('SDG 9 — Industry & Infrastructure', 'Deployment of digital infrastructure in agriculture; promotion of FPO digitisation.'),
    ('SDG 10 — Reduced Inequalities', 'Bridging urban-rural digital divide; vernacular UX for low-literacy users.'),
    ('SDG 12 — Responsible Consumption', 'Traceable supply chain; reduction in chemical-intensive farming through precision advisory.'),
    ('SDG 13 — Climate Action', 'Carbon-credit pilot; promotion of zero-tillage, crop rotation, and bio-inputs.'),
    ('SDG 15 — Life on Land', 'Soil-health monitoring; reduction in pesticide overuse; promotion of bio-fertilizers.'),
])

add_sub_heading('9.3 Ethical Considerations')

add_minor_heading('A. Data Privacy')
add_bullet('Full compliance with the Digital Personal Data Protection Act (DPDPA), 2023.')
add_bullet('Explicit, granular, language-localised consent before collecting farm location, Aadhaar (only when required for KYC), or biometric data.')
add_bullet('All data stored encrypted (AES-256 at rest, TLS 1.3 in transit) within Indian data centres.')
add_bullet('Farmers can download or delete their data at any time via in-app settings (Right to Erasure).')
add_bullet('Annual third-party privacy audits; transparent annual transparency report.')

add_minor_heading('B. Ethical Artificial Intelligence')
add_bullet('Crop-loan credit-scoring algorithms audited quarterly for bias against region, caste, gender, or land size.')
add_bullet('AI advisory clearly labels confidence levels and recommends consulting Krishi Vigyan Kendras for irreversible decisions — never replacing human expertise.')
add_bullet('Pesticide recommendations reviewed by an ICAR-certified agronomy panel before deployment.')
add_bullet('Explainability built into model outputs — every recommendation comes with the data that drove it.')

add_minor_heading('C. Fair Pricing and Anti-Exploitation')
add_bullet('No predatory commission on distress transactions (e.g., farmers losing crop to disease).')
add_bullet('Transparent fee disclosure — every screen shows what KrishiConnect earns from a transaction.')
add_bullet('Buyer rating and verification system ensures bad actors (delayed payments, weight cheating) are rapidly removed.')
add_bullet('No banned or red-label pesticides listed on the platform.')

add_minor_heading('D. Environmental Ethics')
add_bullet('Promotion of bio-fertilizers and Integrated Pest Management.')
add_bullet('Carbon-credit pilot for farmers practicing zero-tillage and crop rotation — generating an additional revenue line for them.')
add_bullet('Logistics partner selection prioritises low-emission options where feasible.')

doc.add_page_break()

# ============================================================
# 10. RISK MANAGEMENT
# ============================================================
add_chapter_heading('10. Risk Management')

add_sub_heading('10.1 Risk Register')
add_minor_heading('Table 10.1 — Risk Register')
add_table([
    ('Category', 'Risk', 'Likelihood / Impact'),
    ('Financial', 'Slow user monetisation; longer-than-expected runway burn', 'High / High'),
    ('Regulatory', 'Sudden changes in APMC laws, MSP regulations, or DPDPA enforcement', 'Medium / High'),
    ('Technical', 'Cybersecurity breach exposing farmer KYC and financial data', 'Medium / Critical'),
    ('Market Adoption', 'Low digital literacy hampering app uptake in deep-rural districts', 'High / Medium'),
    ('Operational', 'Logistics partner failure during peak harvest leading to spoilage and refunds', 'Medium / High'),
    ('Competition', 'Big-tech (Reliance JioKrishi, Amazon Kisan) entering with subsidised pricing', 'Medium / High'),
    ('Climate / Macro', 'Successive bad monsoons reducing transaction volumes', 'Medium / High'),
    ('Talent', 'Key engineering or AI talent attrition to better-funded competitors', 'Medium / Medium'),
    ('Supply-side', 'Counterfeit input infiltration on platform damaging trust', 'Low / High'),
    ('Reputational', 'Misuse of platform by buyers (price collusion, weight manipulation)', 'Medium / Medium'),
])

add_sub_heading('10.2 Mitigation Plan')

add_minor_heading('A. Financial Risk Mitigation')
add_bullet('Maintain minimum 18 months of runway; raise funding in tranches tied to clear KPIs.')
add_bullet('Diversify revenue across 5 streams.')
add_bullet('Apply for non-dilutive grants (NABARD Agri-Tech Fund, ICAR Innovation Grants, Atal Innovation Mission).')

add_minor_heading('B. Regulatory Risk Mitigation')
add_bullet('Onboard a regulatory advisor with experience in APMC and digital-finance laws.')
add_bullet('Active membership in FICCI and NASSCOM Agri-Tech committees.')
add_bullet('Build data-residency and compliance modules from day one.')

add_minor_heading('C. Technical Risk Mitigation')
add_bullet('Annual third-party VAPT (Vulnerability Assessment and Penetration Testing) audits.')
add_bullet('SOC 2 Type II certification target within 24 months.')
add_bullet('Bug-bounty programme; ISO 27001 ISMS framework.')
add_bullet('Real-time anomaly detection on transaction patterns.')

add_minor_heading('D. Market-Adoption Risk Mitigation')
add_bullet('Voice-first, icon-driven UX for low-literacy users; on-ground demos at every village panchayat.')
add_bullet('Partner with FPOs that already enjoy farmer trust.')
add_bullet('Referral incentives to drive word-of-mouth.')

add_minor_heading('E. Operational Risk Mitigation')
add_bullet('Multi-vendor logistics strategy — never more than 35% volume with a single partner.')
add_bullet('Insurance for in-transit produce; SLA-driven contracts with monetary penalties.')
add_bullet('Surge capacity planning before known harvest peaks.')

add_minor_heading('F. Competitive Risk Mitigation')
add_bullet('Build deep moats: proprietary AI models trained on KrishiConnect data, hyperlocal FPO partnerships, and ONDC-first architecture.')
add_bullet('Focus on Tier-3 and Tier-4 agrarian belts that big-tech ignores in initial expansion.')

add_minor_heading('G. Climate / Macro Risk Mitigation')
add_bullet('Diversify across 12+ crop categories and 10+ states.')
add_bullet('Promote crop insurance heavily.')
add_bullet('Maintain 6-month operational reserve.')

add_minor_heading('H. Talent Risk Mitigation')
add_bullet('Competitive ESOP programme — 10 percent equity pool reserved for employees.')
add_bullet('Remote-first culture to widen talent pool.')
add_bullet('Continuous learning budget per employee.')

add_minor_heading('I. Supply-side Risk Mitigation')
add_bullet('Direct manufacturer tie-ups with verified quality certificates.')
add_bullet('Random batch testing at FPO collection centres.')
add_bullet('Strict three-strike policy for counterfeit suppliers.')

doc.add_page_break()

# ============================================================
# 11. CONCLUSION AND FUTURE SCOPE
# ============================================================
add_chapter_heading('11. Conclusion and Future Scope')

add_sub_heading('11.1 Project Summary')
add_body(
    'KrishiConnect represents a meaningful, market-validated, and technology-anchored response '
    'to one of India\'s most enduring developmental challenges — the structural inefficiency '
    'of its agricultural value chain. By unifying market access, input supply, equipment '
    'rental, advisory, and finance in a single, vernacular, AI-driven mobile platform, '
    'KrishiConnect creates value for every stakeholder: farmers earn more, buyers source '
    'cheaper and traceable produce, lenders access a low-default borrower pool, and the '
    'broader economy benefits from reduced food waste and rural distress.'
)
add_body(
    'The startup\'s differentiated USPs — voice-first UX, reverse auctions, AI Crop Doctor, '
    'ONDC-native architecture, and embedded finance — combined with a multi-stream revenue '
    'model and tailwinds from India\'s Digital Agriculture Mission, position it not just as '
    'a viable business but as a potentially category-defining platform.'
)

add_sub_heading('11.2 Five-Year Strategic Roadmap')

add_minor_heading('Year 1 (2026-27) — Pilot and Validation')
add_bullet('Launch in 10 districts of Uttar Pradesh; reach 50,000 active farmers and 1,200 buyers.')
add_bullet('Achieve product-market fit; iterate AI advisor based on field feedback.')
add_bullet('Close pre-seed and seed rounds totalling ₹9 crore.')

add_minor_heading('Year 2 (2027-28) — Geographic Expansion')
add_bullet('Expand to Uttar Pradesh, Madhya Pradesh, Bihar, and Maharashtra; 3 lakh active farmers, 7,500 buyers.')
add_bullet('Launch KrishiPro premium subscription and equipment-rental network.')
add_bullet('Series A funding round (USD 10-12 million).')

add_minor_heading('Year 3 (2028-29) — National Coverage')
add_bullet('Operations in 18 states; 12 lakh farmers, 35,000 buyers; ₹2,400 crore GMV.')
add_bullet('Cross EBITDA break-even.')
add_bullet('Launch B2B export module for traceable Indian produce (basmati, spices, organic).')

add_minor_heading('Year 4 (2029-30) — Vertical Deepening')
add_bullet('Add dairy, poultry, fisheries, and apiculture marketplaces.')
add_bullet('Launch carbon-credit monetisation for sustainable farmers.')
add_bullet('Series B funding (USD 30-50 million); aim for 25 million registered farmers.')

add_minor_heading('Year 5 (2030-31) — International and IPO Readiness')
add_bullet('Pilot in Bangladesh, Nepal, Sri Lanka, and select African countries (Kenya, Nigeria) where smallholder dynamics mirror India.')
add_bullet('Strategic partnerships with FAO, World Bank IFC, and ADB for co-funded rural digitisation.')
add_bullet('IPO readiness via DRHP filing — targeting valuation of USD 1.2 to 1.5 billion (Unicorn status).')

add_sub_heading('11.3 Long-Term Future Scope')
add_bullet('Vertical Expansion — fisheries, dairy, apiculture, and agri-tourism.')
add_bullet('Geographic Expansion — Bangladesh, Nepal, Sri Lanka, Kenya, Nigeria.')
add_bullet('New Revenue Lines — carbon credits, agri-data licensing to research institutions, agri-input financing for FPOs.')
add_bullet('Deeper AI — generative models for personalised crop calendars, climate-resilient variety recommendation, and multi-modal advisory (image + voice + sensor).')
add_bullet('Web3 and Tokenisation — fractional ownership of high-value equipment, tokenised carbon credits, and decentralised farmer cooperatives.')

add_sub_heading('11.4 Closing Reflection')
add_body(
    'The greatest strength of KrishiConnect lies in its alignment of profit and purpose — '
    'every rupee the platform earns translates into measurable benefit at the farm gate. In '
    'an era where technology has reshaped finance, retail, mobility, and entertainment, '
    'agriculture remains the last great frontier of digital transformation. KrishiConnect is '
    'conceived to lead that transformation in India — not by displacing the farmer, but by '
    'amplifying his or her voice, bargaining power, and future.'
)
add_body(
    'This Startup and Entrepreneurial Activity Report has demonstrated, through structured '
    'analysis, that the proposed venture is technologically feasible, commercially viable, '
    'socially impactful, and strategically timely. The journey from an academic concept to '
    'a fully scaled enterprise is, of course, long and uncertain — but the convergence of '
    'public infrastructure, market readiness, and capital availability suggests that the '
    'window for action is open now. The work done in compiling this report has, for the '
    'author, served not merely as an academic exercise but as the first concrete step in a '
    'longer entrepreneurial journey.'
)

doc.add_page_break()

# ============================================================
# 12. REFERENCES
# ============================================================
add_chapter_heading('12. References')

refs = [
    'Ministry of Agriculture & Farmers Welfare, Government of India. (2025). Annual Report 2024-25. Retrieved from https://agriculture.gov.in.',
    'NABARD. (2023). All-India Rural Financial Inclusion Survey (NAFIS) 2021-22. National Bank for Agriculture and Rural Development, Mumbai.',
    'Bain & Company and EY-NASSCOM. (2024). Indian Agritech: Unlocking the USD 34 Billion Opportunity. Joint Industry Report.',
    'NITI Aayog. (2024). Doubling Farmers\' Income — Implementation Strategy. Government of India.',
    'Open Network for Digital Commerce (ONDC). (2025). ONDC for Agriculture — Reference Architecture v2.1. Retrieved from https://ondc.org.',
    'Ministry of Electronics and Information Technology. (2023). The Digital Personal Data Protection Act, 2023. Government of India.',
    'ICAR. (2024). Vision 2050 — Indian Council of Agricultural Research. New Delhi.',
    'World Bank. (2023). India Agriculture Overview. Washington D.C.',
    'FAO. (2024). The State of Food and Agriculture 2024. Food and Agriculture Organization of the United Nations.',
    'IBEF. (2025). Indian Agriculture and Allied Industries Report. India Brand Equity Foundation.',
    'TechSci Research. (2024). India Agritech Market — Forecast and Opportunities, 2030.',
    'McKinsey & Company. (2023). Harvesting Golden Opportunities in Indian Agriculture.',
    'Inc42. (2025). Indian Agritech Funding Report — H1 2025. Retrieved from https://inc42.com.',
    'Ministry of Rural Development. (2025). Self-Help Group and FPO Implementation Status. Government of India.',
    'Kshetri, N. (2022). Blockchain and Sustainable Supply Chain Management in Developing Countries. International Journal of Information Management, 60.',
    'DeHaat. (2024). Annual Impact Report. Retrieved from https://agrevolution.in.',
    'AgroStar. (2024). Farmer Engagement Insights Report. Retrieved from https://agrostar.in.',
    'Press Information Bureau. (2024). Digital Agriculture Mission — Cabinet Approval Note. Retrieved from https://pib.gov.in.',
    'IIM Ahmedabad — CMA. (2023). Smallholder Farmer Income and Access Study. Centre for Management in Agriculture.',
    'KANTAR. (2025). Rural Smartphone Index — KANTAR ICUBE 2025 Report.',
    'Reserve Bank of India. (2024). Report on Trend and Progress of Banking in India 2023-24. RBI Mumbai.',
    'PwC India. (2024). Indian Agritech: A USD 24 Billion Opportunity. PricewaterhouseCoopers India.',
    'CII. (2024). Vision 2030 for Indian Agriculture. Confederation of Indian Industry.',
    'Department of Agricultural Research and Education. (2024). Annual Report 2023-24. Ministry of Agriculture and Farmers\' Welfare, GoI.',
    'World Economic Forum. (2024). Innovation with a Purpose: Improving Agricultural Value Chains. Geneva: WEF.',
]

for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'[{i}] {r}')
    set_run(run, size=12)

doc.add_page_break()

# ============================================================
# 13. APPENDICES
# ============================================================
add_chapter_heading('13. Appendices')

add_sub_heading('Appendix A — Glossary of Terms')
glossary = [
    ('AgriStack', 'Government-of-India digital framework for unified farmer database and services.'),
    ('APMC', 'Agricultural Produce Market Committee — state-regulated wholesale produce markets.'),
    ('CAC', 'Customer Acquisition Cost — average cost incurred to acquire one paying user.'),
    ('CAGR', 'Compound Annual Growth Rate.'),
    ('DPDPA', 'Digital Personal Data Protection Act, 2023 — India\'s primary data-privacy law.'),
    ('e-NAM', 'Electronic National Agriculture Market — government-operated online trading portal.'),
    ('EBITDA', 'Earnings Before Interest, Taxes, Depreciation and Amortisation.'),
    ('FPO', 'Farmer Producer Organisation — collective of farmers registered as a producer company.'),
    ('GMV', 'Gross Merchandise Value — total value of goods transacted on the platform.'),
    ('ICAR', 'Indian Council of Agricultural Research — apex agricultural research body.'),
    ('KVK', 'Krishi Vigyan Kendra — district-level agricultural extension centre.'),
    ('LSTM', 'Long Short-Term Memory — a recurrent neural-network architecture used for time-series prediction.'),
    ('MSP', 'Minimum Support Price — government-guaranteed floor price for select crops.'),
    ('NDVI', 'Normalised Difference Vegetation Index — satellite-derived measure of crop health.'),
    ('ONDC', 'Open Network for Digital Commerce — India\'s open commerce protocol.'),
    ('PMFBY', 'Pradhan Mantri Fasal Bima Yojana — flagship crop-insurance scheme.'),
    ('SOM / SAM / TAM', 'Serviceable Obtainable / Serviceable Addressable / Total Addressable Market.'),
    ('TSP', 'Technology Service Provider — entity offering ONDC integration services.'),
]
add_table([('Term', 'Definition')] + glossary, header=True, widths=[1.5, 4.6])

add_sub_heading('Appendix B — Sample Wireframe Description')
add_body(
    'The KrishiConnect mobile application is structured around a five-tab bottom navigation '
    'bar — Home, Sell, Buy Inputs, Equipment, and Advisory. The Home screen shows '
    'personalised mandi prices for the farmer\'s declared crops, weather forecasts, urgent '
    'pest alerts (if any), and a prominent "Ask KrishiSakhi" voice button. The Sell flow '
    'allows the farmer to list produce in three steps: select crop, enter quantity and grade '
    '(with on-device camera grading assistance), and accept the highest reverse-auction bid '
    'within a 24-hour window. The Buy Inputs flow features curated input bundles by crop and '
    'season, with full ingredient transparency and verified-supplier badges. The Equipment '
    'flow shows nearby tractors, harvesters, and drones available for hire, with hourly '
    'pricing and instant booking. The Advisory flow combines the AI Crop Doctor (photo-based '
    'diagnosis), KrishiSakhi voice chatbot, and crop-specific calendars synced to the user\'s '
    'sowing dates.'
)

add_sub_heading('Appendix C — Sample Daily Usage Scenario')
add_body(
    'Ramesh, a 42-year-old smallholder in Bundelkhand, opens the KrishiConnect app at 7 AM. '
    'The home screen greets him in Bundeli with today\'s mandi prices for wheat (₹2,310 per '
    'quintal at Hamirpur, ₹2,365 at Mahoba), a weather alert about expected showers tomorrow, '
    'and a reminder that his pigeon-pea field is 78 days into cultivation — the optimal '
    'window for the second irrigation. He taps the AI Crop Doctor and clicks a photo of a '
    'yellowing leaf; the model identifies early-stage Yellow Mosaic Virus and recommends a '
    'specific bio-pesticide that he can order from the in-app store, with delivery in 36 '
    'hours. In the evening, he posts 12 quintals of harvested wheat for sale; within four '
    'hours, three buyers have placed bids, the highest being ₹2,418 per quintal — '
    '₹108 above the prevailing mandi rate. He accepts the bid; payment is escrowed, the '
    'logistics partner is notified, and the wheat is collected the next morning. Ramesh '
    'receives the payment within 24 hours of pickup, with all fees clearly disclosed.'
)

add_sub_heading('Appendix D — Founder\'s Note')
add_body(
    'When I began researching the topic of an Agri-Tech marketplace, I was struck by a '
    'simple but jarring statistic: the cultivator who grows our food earns less per hour '
    'than the delivery rider who brings it to our door. The injustice of that imbalance '
    'animates this report. While the work documented here is academic, my hope is that '
    'the ideas — voice-first vernacular UX, reverse-auction pricing, embedded finance, and '
    'AI-driven advisory — find their way into a real venture, whether built by me or by '
    'someone reading this report. The Indian farmer has waited long enough.'
)
add_body(' ')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(f'— {STUDENT_NAME}\nMCA, 2025-26'); set_run(r, size=12, italic=True)

doc.add_page_break()

# ============================================================
# SIGNATURES
# ============================================================
add_chapter_heading('Signatures', text=None) if False else None  # placeholder
add_sub_heading('Signatures')
add_body(' '); add_body(' '); add_body(' ')

t = doc.add_table(rows=2, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = '_________________________'
t.rows[0].cells[1].text = '_________________________'
t.rows[1].cells[0].text = f'{STUDENT_NAME}\n(Student)\nRoll No: {ROLL_NO}\n{BATCH}'
t.rows[1].cells[1].text = f'{FACULTY}\n(Subject Incharge / Supervisor)\nAssistant Professor\nComputer Application Department'
for r in t.rows:
    for c in r.cells:
        for para in c.paragraphs:
            for run in para.runs:
                set_run(run, size=12)

doc.save(OUT)
print(f"Document generated: {OUT}")
